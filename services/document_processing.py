"""Document ingestion utilities supporting PDF, DOCX, and TXT."""
from __future__ import annotations

from pathlib import Path
from typing import Dict

import docx
import PyPDF2


def ingest(path: str) -> Dict[str, str]:
    """Return document text and metadata from the given file path."""
    p = Path(path)
    if p.suffix.lower() == ".pdf":
        with p.open("rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        meta = reader.metadata or {}
    elif p.suffix.lower() in {".doc", ".docx"}:
        doc = docx.Document(p)
        text = "\n".join(p.text for p in doc.paragraphs)
        meta = {"paragraphs": len(doc.paragraphs)}
    else:
        text = p.read_text()
        meta = {}
    return {"text": text, "metadata": meta}
